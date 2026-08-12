"""
File creation tools — write documents/PDFs and save them to the desktop (or elsewhere).
"""
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

from core.config import Config
from core.safety import desktop_destination


def _resolve_path(filename: str, ext: str) -> Path:
    return desktop_destination(Config.DESKTOP_PATH, filename, ext)


def create_text_file(filename: str, content: str) -> str:
    """Create a plain .txt file on the desktop with the given content."""
    path = _resolve_path(filename, ".txt")
    path.write_text(content, encoding="utf-8")
    return f"Saved text file to {path}"


def create_word_document(filename: str, title: str, body: str) -> str:
    """Create a Word (.docx) document with a title and body text, saved to the desktop."""
    path = _resolve_path(filename, ".docx")
    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for paragraph in body.split("\n\n"):
        doc.add_paragraph(paragraph)
    doc.save(str(path))
    return f"Saved Word document to {path}"


def create_pdf(filename: str, title: str, body: str) -> str:
    """Create a simple PDF with a title and wrapped body text, saved to the desktop."""
    path = _resolve_path(filename, ".pdf")
    c = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    y = height - inch

    if title:
        c.setFont("Helvetica-Bold", 16)
        c.drawString(inch, y, title)
        y -= 0.4 * inch

    c.setFont("Helvetica", 11)
    max_chars_per_line = 95
    for paragraph in body.split("\n"):
        words = paragraph.split(" ")
        line = ""
        for word in words:
            if len(line) + len(word) + 1 > max_chars_per_line:
                c.drawString(inch, y, line)
                y -= 0.22 * inch
                line = word
                if y < inch:
                    c.showPage()
                    c.setFont("Helvetica", 11)
                    y = height - inch
            else:
                line = f"{line} {word}".strip()
        if line:
            c.drawString(inch, y, line)
            y -= 0.22 * inch
        y -= 0.1 * inch

    c.save()
    return f"Saved PDF to {path}"
